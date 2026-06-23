# 负责读取pdf，docx, txt,md等文档

from pathlib import Path
from typing import List, Tuple
from langchain_core.documents import Document

SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".md",
}


def load_document_file(file_path: str | Path) -> List[Document]:
    """
    根据文件后缀读取文档，并统一转换为 LangChain Document。

    支持：
    - PDF
    - DOCX
    - TXT
    - MD
    """
    path = Path(file_path) 
    suffix = path.suffix.lower() 
    
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"不支持的文件格式: {suffix}，当前仅支持 PDF / DOCX / TXT / MD"
        )
    
    if suffix == ".pdf":
        return _load_pdf(path)

    if suffix == ".docx":
        return _load_docx(path)

    if suffix in {".txt", ".md"}:
        return _load_text(path)

    raise ValueError(f"无法处理文件: {path}")    
    

def _base_metadata(path: Path) -> dict:
    return {
        "source": str(path),
        "file_name": path.name,
        "file_ext": path.suffix.lower(),
        "source_type": "uploaded_file",
    }
    

def _load_text(path: Path) -> List[Document]:
    """
    读取TXT / MD 
    """
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = path.read_text(encoding="gbk", errors="ignore") 
    
    if not text.strip():
        return []  
    
    return [
        Document(
            page_content=text,
            metadata=_base_metadata(path)
        )
    ]   
    
    
def _load_pdf(path: Path) -> List[Document]:
    """
    读取 PDF。

    注意：
    这个方法适合文本型 PDF。
    如果是扫描版 PDF，需要后续接 OCR。
    """
    from pypdf import PdfReader
    
    reader = PdfReader(str(path)) 
    documents: List[Document] = []     
    
    for page_index, page in enumerate(reader.pages):
        text = page.extract_text() or "" 
        
        if not text.strip():
            continue
        
        metadata = _base_metadata(path)
        metadata.update(
            {
                "page": page_index + 1,
            }
        )
        
        documents.append(
            Document(
                page_content=text,
                metadata=metadata
            )
        )
    
    return documents  


def _load_docx(path: Path) -> List[Document]:
    """
    读取 DOCX。
    """ 
    from docx import Document as DocxDocument
    
    docx = DocxDocument(str(path))
    
    paragraphs = [
        p.text.strip()
        for p in docx.paragraphs
        if p.text and p.text.strip()
    ] 
    
    text = "\n".join(paragraphs) 
    
    if not text.strip():
        return []
    
    return [
        Document(
            page_content=text,
            metadata=_base_metadata(path)
        )
    ]

      